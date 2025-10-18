from flask import Flask, request, jsonify, render_template, make_response, Response, stream_with_context
from werkzeug.exceptions import HTTPException
import os
import logging
from logging.handlers import RotatingFileHandler
import threading
import uuid
import io
import time
import hmac
import hashlib
import json as _json
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_cors import CORS
from typing import List, Dict, Set
from concurrent.futures import ThreadPoolExecutor, as_completed
from utils.text import normalize_text as _normalize_text
from utils.text import tokenize_fa as _tokenize_fa
from utils.text import split_paragraphs as _split_paragraphs
from utils.text import extract_text_from_html
from services.learner import VectorLearner

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": os.getenv('CORS_ORIGINS', '*')}})
limiter = Limiter(get_remote_address, app=app, default_limits=[os.getenv('RATE_LIMIT_DEFAULT', '60 per minute')])

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(BASE_DIR, 'config.json')

# پوشهٔ پیش‌فرض PDFها (خواندن خودکار)
DEFAULT_PDF_DIR = os.path.join(BASE_DIR, 'pdf')
VECTOR_INDEX_DIR = os.path.join(BASE_DIR, 'vector_index')
SOURCES_DIR = os.path.join(BASE_DIR, 'sources')

# تنظیم سادهٔ لاگ‌ها
def _configure_logging():
    level_name = os.getenv('LOG_LEVEL', 'INFO').upper()
    level = getattr(logging, level_name, logging.INFO)
    logging.basicConfig(level=level, format='[%(asctime)s] %(levelname)s %(message)s')

_configure_logging()
logger = logging.getLogger("app")

# ساخت لاگ JSON با چرخش فایل
def _setup_json_logging():
    try:
        logs_dir = os.path.join(BASE_DIR, 'logs')
        os.makedirs(logs_dir, exist_ok=True)
        file_path = os.path.join(logs_dir, 'app.log')

        class JsonFormatter(logging.Formatter):
            def format(self, record: logging.LogRecord) -> str:
                payload = {
                    'ts': int(time.time()),
                    'level': record.levelname,
                    'msg': record.getMessage(),
                    'name': record.name,
                }
                if hasattr(record, 'extra') and isinstance(record.extra, dict):
                    payload.update(record.extra)
                return _json.dumps(payload, ensure_ascii=False)

        handler = RotatingFileHandler(file_path, maxBytes=2 * 1024 * 1024, backupCount=5, encoding='utf-8')
        handler.setFormatter(JsonFormatter())
        root = logging.getLogger()
        root.addHandler(handler)
    except Exception:
        pass

_setup_json_logging()

# درخواست‌نگار ساده: زمان پاسخ را در لاگ ثبت می‌کند
@app.before_request
def _before_request_timer():
    try:
        setattr(request, '_start_time', time.perf_counter())
    except Exception:
        pass

@app.after_request
def _after_request_timer(response):
    try:
        start = getattr(request, '_start_time', None)
        if start is not None:
            dur_ms = int((time.perf_counter() - start) * 1000)
            logger.info("%s %s %s %dms", request.method, request.path, str(response.status_code), dur_ms)
    except Exception:
        pass
    return response

# آپلود حذف شده است؛ نیازی به الگوی پسوندها نیست

# ----------- ذخیرهٔ متن‌ها در حافظه بر اساس شناسهٔ کاربر (کوکی) -----------
SESSION_CONTEXTS: Dict[str, List[str]] = {}
_context_lock = threading.Lock()
SESSION_PROCESSED: Dict[str, Set[str]] = {}
TOTAL_TEXTS_COUNT: int = 0
PROCESSED_FILES_COUNT: int = 0
PREWARM_PROGRESS: Dict[str, Dict[str, object]] = {}

# نمایهٔ سبک برای بازیابی سریع پاراگراف‌ها (بدون وابستگی خارجی)
# per-session index: client_id -> { 'paragraphs': List[str], 'inverted': Dict[str, Set[int]] }
SESSION_INDEX: Dict[str, Dict[str, object]] = {}

# نمایهٔ向向向向向 (FAISS) سراسری: برای «یادگیری» قوانین و بازیابی برداری
GLOBAL_VECTOR = {
    'index': None,           # faiss.Index یا None
    'paragraphs': [],        # List[str]
    'model_name': None,      # نام مدل تع嵌ه
}
EMBEDDER = { 'model': None, 'name': None }
LEARNER = VectorLearner(BASE_DIR)

# پاسخ‌های بازبینی‌شدهٔ انسانی (Human-in-the-loop)
CURATED_ANSWERS: Dict[str, Dict[str, object]] = {}

# وضعیت آخرین تماس مدل
MODEL_STATUS: Dict[str, object] = {
    'engine': 'local',
    'error': '',
    'latency_ms': 0,
    'ts': 0,
}

def _read_config_file() -> Dict[str, object]:
    try:
        import json as _json
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
                return _json.load(f) or {}
    except Exception:
        pass
    return {}

def _write_config_file(cfg: Dict[str, object]) -> bool:
    try:
        import json as _json
        with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
            _json.dump(cfg or {}, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

def _mask_secret(value: str) -> str:
    if not value:
        return ''
    if len(value) <= 6:
        return '*' * len(value)
    return ('*' * (len(value) - 4)) + value[-4:]

def _load_persisted_config_into_env() -> None:
    cfg = _read_config_file()
    # DeepSeek
    dk = str(cfg.get('DEEPSEEK_API_KEY') or '').strip()
    dm = str(cfg.get('DEEPSEEK_MODEL') or '').strip()
    if dk:
        os.environ['DEEPSEEK_API_KEY'] = dk
        os.environ['USE_DEEPSEEK'] = '1'
    if dm:
        os.environ['DEEPSEEK_MODEL'] = dm
    # Optional toggles
    if str(cfg.get('USE_OLLAMA') or ''):
        os.environ['USE_OLLAMA'] = str(cfg.get('USE_OLLAMA'))
    if str(cfg.get('USE_LLAMA') or ''):
        os.environ['USE_LLAMA'] = str(cfg.get('USE_LLAMA'))

# بارگذاری تنظیمات ذخیره‌شده هنگام راه‌اندازی فرآیند
_load_persisted_config_into_env()

def _get_or_create_client_id() -> str:
    client_id = request.cookies.get('client_id')
    if client_id and isinstance(client_id, str) and len(client_id) >= 8:
        return client_id
    return uuid.uuid4().hex

def _is_fast_mode() -> bool:
    # حالت سریع برای پاسخ‌دهی کوتاه‌تر و محدودسازی زمینه
    return _str_to_bool(os.getenv('FAST_MODE', '1'), default=True)

# متن: به utils/text منتقل شد

def _ensure_dirs():
    try:
        os.makedirs(VECTOR_INDEX_DIR, exist_ok=True)
    except Exception:
        pass
    try:
        os.makedirs(SOURCES_DIR, exist_ok=True)
    except Exception:
        pass

def _get_embedder():
    # cached embedder
    if EMBEDDER.get('model') is not None:
        return EMBEDDER['model'], EMBEDDER.get('name')
    try:
        from sentence_transformers import SentenceTransformer
        model_name = os.getenv('EMBED_MODEL', 'paraphrase-multilingual-MiniLM-L12-v2')
        model = SentenceTransformer(model_name)
        EMBEDDER['model'] = model
        EMBEDDER['name'] = model_name
        return model, model_name
    except Exception:
        return None, None

# استخراج متن HTML: به utils/text منتقل شد

def _faiss_available():
    try:
        import faiss  # type: ignore
        _ = faiss
        return True
    except Exception:
        return False

def build_global_vector_index(sources: List[str] | None = None) -> Dict[str, object]:
    """ساخت ایندکس برداری با VectorLearner از سشن‌ها و پوشهٔ پیش‌فرض."""
    _ensure_dirs()
    texts: List[str] = []
    with _context_lock:
        for arr in (SESSION_CONTEXTS.values() or []):
            for t in (arr or []):
                texts.append(t)
    srcs = sources or [DEFAULT_PDF_DIR]
    for src in srcs:
        tlist, _previews, _cnt = ingest_pdfs_from_dir(src, recursive=True, client_id=None, fast=True)
        texts.extend(tlist)
    res = LEARNER.build_from_texts(texts)
    # همگام‌سازی وضعیت قدیمی برای سازگاری با پنل
    if res.get('built'):
        GLOBAL_VECTOR['index'] = LEARNER.index
        GLOBAL_VECTOR['paragraphs'] = LEARNER.paragraphs
        GLOBAL_VECTOR['model_name'] = LEARNER.model_name
    return res

def _load_global_vector_index_if_exists():
    if GLOBAL_VECTOR.get('index') is not None:
        return True
    if LEARNER.load_if_exists():
        GLOBAL_VECTOR['index'] = LEARNER.index
        GLOBAL_VECTOR['paragraphs'] = LEARNER.paragraphs
        GLOBAL_VECTOR['model_name'] = LEARNER.model_name
        return True
    return False

def _retrieve_with_vector(query: str, top_k: int = 5) -> List[Dict[str, object]]:
    if not query:
        return []
    if LEARNER.index is None and not _load_global_vector_index_if_exists():
        return []
    return LEARNER.search(query, top_k=top_k)

def _ensure_session_index(client_id: str) -> None:
    if not client_id:
        return
    with _context_lock:
        idx = SESSION_INDEX.get(client_id)
        texts = SESSION_CONTEXTS.get(client_id) or []
    if idx and idx.get('built_from_count') == len(texts):
        return
    paragraphs: List[str] = []
    for t in texts:
        paragraphs.extend(_split_paragraphs(t))
    inverted: Dict[str, Set[int]] = {}
    for i, p in enumerate(paragraphs):
        for tok in set(_tokenize_fa(p)):
            s = inverted.get(tok)
            if s is None:
                s = set()
                inverted[tok] = s
            s.add(i)
    with _context_lock:
        SESSION_INDEX[client_id] = {
            'paragraphs': paragraphs,
            'inverted': inverted,
            'built_from_count': len(texts),
        }

def _retrieve_paragraphs(question: str, client_id: str, top_k: int = 5) -> List[Dict[str, object]]:
    # ابتدا تلاش با ایندکس برداری سراسری (اگر موجود)
    vec = _retrieve_with_vector(question, top_k=top_k)
    if vec:
        return vec
    # سپس ایندکس محلی سشن
    if not client_id:
        return []
    with _context_lock:
        idx = SESSION_INDEX.get(client_id)
    if not idx:
        return []
    paragraphs: List[str] = idx.get('paragraphs') or []
    inverted: Dict[str, Set[int]] = idx.get('inverted') or {}
    q_tokens = _tokenize_fa(question)
    if not q_tokens:
        return []
    from collections import Counter
    scores = Counter()
    for qt in set(q_tokens):
        for pid in inverted.get(qt, set()):
            scores[pid] += 1
    if not scores:
        return []
    ranked = scores.most_common(top_k)
    results: List[Dict[str, object]] = []
    for pid, sc in ranked:
        snippet = paragraphs[pid]
        results.append({'pid': pid, 'score': int(sc), 'snippet': snippet})
    return results

def _build_limited_context(texts: List[str], per_text_limit: int = 1200, total_limit: int = 9000) -> str:
    # برش هر متن و همچنین محدودیت حداکثر کل کاراکترها
    if not texts:
        return ''
    clipped_parts: List[str] = []
    remaining = max(1000, int(total_limit))
    per_limit = max(300, int(per_text_limit))
    for t in texts:
        if remaining <= 0:
            break
        part = (t or '')[:per_limit]
        if not part:
            continue
        if len(part) > remaining:
            part = part[:remaining]
        clipped_parts.append(part)
        remaining -= len(part)
    return "\n\n---\n\n".join(clipped_parts)

def _append_session_context(client_id: str, texts: List[str]) -> None:
    added = [t for t in texts if t]
    if not added:
        return
    global TOTAL_TEXTS_COUNT
    with _context_lock:
        existing = SESSION_CONTEXTS.get(client_id) or []
        existing.extend(added)
        SESSION_CONTEXTS[client_id] = existing
        TOTAL_TEXTS_COUNT += len(added)

# ----------- کمکی: خواندن PDFها از یک پوشه -----------
def ingest_pdfs_from_dir(dir_path: str, recursive: bool = True, max_files: int = 200, client_id: str | None = None, fast: bool = False):
    dir_path = os.path.expanduser(dir_path)
    if not os.path.isabs(dir_path):
        dir_path = os.path.abspath(os.path.join(BASE_DIR, dir_path))
    if not os.path.isdir(dir_path):
        return [], [], 0

    def iter_pdfs(path: str):
        if recursive:
            for root, _dirs, files in os.walk(path):
                for name in files:
                    if name.lower().endswith('.pdf'):
                        yield os.path.join(root, name)
        else:
            for name in os.listdir(path):
                fp = os.path.join(path, name)
                if os.path.isfile(fp) and name.lower().endswith('.pdf'):
                    yield fp

    saved_texts: List[str] = []
    previews: List[str] = []
    total_files = 0
    errors_count = 0
    start_ts = time.time()
    try:
        max_seconds = float(os.getenv('INGEST_MAX_SECONDS', '30'))
    except Exception:
        max_seconds = 30.0
    # مسیرهای قبلاً پردازش‌شده برای این سشن
    processed: Set[str] = set()
    if client_id:
        with _context_lock:
            processed = set(SESSION_PROCESSED.get(client_id) or set())
    new_processed: Set[str] = set()
    logger.info("Ingest start: %s (recursive=%s)", dir_path, recursive)
    # گردآوری یک بچ فایل تا سقف max_files
    batch_paths: List[str] = []
    for fp in iter_pdfs(dir_path):
        abs_fp = os.path.abspath(fp)
        if abs_fp in processed:
            continue
        batch_paths.append(abs_fp)
        if len(batch_paths) >= max_files:
            break

    def _extract(path_abs: str) -> tuple[str, str, str]:
        try:
            text = extract_text_fast(path_abs) if fast else extract_text(path_abs)
        except Exception:
            logger.exception("Extract failed: %s", path_abs)
            nonlocal errors_count
            errors_count += 1
            text = ''
        if text:
            pv = (text[:500] + '...') if len(text) > 500 else text
            return (path_abs, text, pv)
        return (path_abs, '', '')

    if batch_paths:
        # اجرای موازی
        try:
            workers = int(os.getenv('INGEST_WORKERS', str(max(2, min(8, (os.cpu_count() or 4))))))
        except Exception:
            workers = max(2, min(8, (os.cpu_count() or 4)))
        try:
            with ThreadPoolExecutor(max_workers=workers) as pool:
                futures = [pool.submit(_extract, p) for p in batch_paths]
                for fut in as_completed(futures):
                    path_abs, text, pv = fut.result()
                    if text:
                        saved_texts.append(text)
                        previews.append(pv)
                        new_processed.add(path_abs)
                        total_files += 1
                    if (time.time() - start_ts) > max_seconds:
                        logger.warning("Ingest time budget exceeded (%ss)", max_seconds)
                        break
        except Exception:
            # fallback سریالی
            for p in batch_paths:
                path_abs, text, pv = _extract(p)
                if text:
                    saved_texts.append(text)
                    previews.append(pv)
                    new_processed.add(path_abs)
                    total_files += 1
                if (time.time() - start_ts) > max_seconds:
                    logger.warning("Ingest time budget exceeded (serial) (%ss)", max_seconds)
                    break

    if client_id and new_processed:
        with _context_lock:
            global PROCESSED_FILES_COUNT
            already = SESSION_PROCESSED.get(client_id) or set()
            already.update(new_processed)
            SESSION_PROCESSED[client_id] = already
            PROCESSED_FILES_COUNT += len(new_processed)
    logger.info("Ingest done: files=%d errors=%d elapsed_ms=%d", total_files, errors_count, int((time.time()-start_ts)*1000))
    return saved_texts, previews, total_files

# ----------- استخراج متن از PDF -----------
def extract_text_from_pdf(pdf_path):
    # import تنبل برای اینکه اگر PyMuPDF نصب نباشد، برنامه حداقل بالا بیاید
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError("کتابخانه PyMuPDF (بسته pymupdf) نصب نیست. دستور نصب: pip install pymupdf") from exc

    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text")
    return text.strip()

# تلاش دوم: استخراج با pdfminer.six (برای برخی PDFهای متنی که PyMuPDF متن نمی‌دهد)
def extract_text_with_pdfminer(pdf_path):
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
    except Exception:
        return ""
    try:
        mined = pdfminer_extract_text(pdf_path)
        return (mined or "").strip()
    except Exception:
        return ""

def extract_text_via_ocr(pdf_path):
    # OCR برای PDFهای اسکن‌شده (در صورت نبود متن)
    try:
        import fitz  # برای رندر صفحات PDF به تصویر
        from PIL import Image  # pillow
        import pytesseract
    except Exception as exc:
        raise RuntimeError("برای OCR نیاز به نصب pillow و pytesseract و Tesseract OCR دارید.") from exc

    # تلاش برای تعیین خودکار مسیر tesseract.exe در ویندوز
    if os.name == 'nt':
        # اگر کاربر متغیر محیطی مشخص کرده باشد، مقدم است
        user_path = os.getenv('TESSERACT_PATH')
        candidates = []
        if user_path:
            candidates.append(user_path)
        # مسیرهای معمول نصب
        candidates.extend([
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
        ])
        for p in candidates:
            if os.path.exists(p):
                try:
                    pytesseract.pytesseract.tesseract_cmd = p
                    break
                except Exception:
                    pass

    doc = fitz.open(pdf_path)
    ocr_text_parts = []
    # مقیاس 3x برای کیفیت بهتر OCR (~216dpi اگر پایه 72dpi باشد)
    upscale_matrix = fitz.Matrix(3, 3)
    for page in doc:
        pix = page.get_pixmap(matrix=upscale_matrix)
        mode = "RGB" if pix.alpha == 0 else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        # تلاش با فارسی+انگلیسی، در صورت نبود داده زبانی، fallback به انگلیسی
        try:
            text_page = pytesseract.image_to_string(img, lang="fas+eng")
        except Exception:
            text_page = pytesseract.image_to_string(img, lang="eng")
        if text_page:
            ocr_text_parts.append(text_page)
    return "\n".join(ocr_text_parts).strip()

def extract_text(pdf_path):
    # ابتدا تلاش برای متن دیجیتال؛ در صورت نبود PyMuPDF یا بروز خطا، ادامه می‌دهیم
    try:
        base_text = extract_text_from_pdf(pdf_path)
    except Exception:
        base_text = ""
    if base_text and len(base_text.strip()) >= 5:
        return base_text.strip()
    # تلاش دوم با pdfminer
    miner_text = extract_text_with_pdfminer(pdf_path)
    if miner_text and len(miner_text.strip()) >= 5:
        return miner_text.strip()
    # در صورت فعال بودن OCR و نبود متن مناسب، OCR انجام می‌شود
    if _str_to_bool(os.getenv('OCR', '1'), default=True):
        try:
            return extract_text_via_ocr(pdf_path)
        except Exception:
            return (base_text or "").strip()
    return (base_text or "").strip()

def extract_text_fast(pdf_path):
    """استخراج سریع: بدون OCR، تلاش سریع با PyMuPDF سپس pdfminer."""
    # تلاش سریع با PyMuPDF
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text("text")
        if text and len(text.strip()) >= 3:
            return text.strip()
    except Exception:
        pass
    # تلاش دوم با pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        mined = pdfminer_extract_text(pdf_path)
        return (mined or "").strip()
    except Exception:
        return ""

# ----------- LLaMA (llama-cpp) یکپارچه‌سازی تنبل -----------
llama_instance = None
llama_error_message = None
_llama_lock = threading.Lock()

def _str_to_bool(value, default=False):
    if value is None:
        return default
    return str(value).strip().lower() in ('1', 'true', 'yes', 'on', 'y')

def _should_use_llama() -> bool:
    # به‌صورت پیش‌فرض غیرفعال تا تمرکز روی Ollama باشد
    return _str_to_bool(os.getenv('USE_LLAMA', '0'), default=False)

def _should_use_ollama() -> bool:
    return _str_to_bool(os.getenv('USE_OLLAMA', '1'), default=True)

def _normalize_text(s: str) -> str:
    if not s:
        return ''
    try:
        import re
        s = s.lower()
        s = re.sub(r"[\u200c\u200f\u200e\ufeff]", "", s)
        s = re.sub(r"[^\w\s\u0600-\u06FF]", " ", s)
        s = re.sub(r"\s+", " ", s)
        return s.strip()
    except Exception:
        return s

def _hash_key_for_question(question: str) -> str:
    try:
        import hashlib
        base = _normalize_text(question)
        return hashlib.sha256(base.encode('utf-8', errors='ignore')).hexdigest()
    except Exception:
        return question[:64]

def _score_paragraph(q_norm: str, p_norm: str) -> int:
    if not q_norm or not p_norm:
        return 0
    q_terms = set(q_norm.split())
    p_terms = set(p_norm.split())
    return len(q_terms & p_terms)

def generate_answer_locally(question: str, context: str, k: int = 3, per_snippet_chars: int = 400) -> str:
    qn = _normalize_text(question)
    if not context:
        return "برای پاسخ دقیق‌تر، لطفاً فایل‌ها را بارگذاری یا زمینه را ارسال کنید."
    parts = [p for p in context.split("\n\n") if p and p.strip()]
    scored = []
    for p in parts:
        pn = _normalize_text(p)
        score = _score_paragraph(qn, pn)
        if score > 0:
            scored.append((score, p))
    if not scored:
        hint = (context[: per_snippet_chars * k] + '...') if len(context) > per_snippet_chars * k else context
        return "پاسخ محلی (بدون مدل):\n" + hint
    scored.sort(key=lambda x: x[0], reverse=True)
    picks = [p for _s, p in scored[:k]]
    clipped = []
    for p in picks:
        clipped.append(p[:per_snippet_chars])
    result = "\n\n---\n\n".join(clipped)
    return "نتیجهٔ سریع بر اساس متن:\n" + result

# ----------- تولید لایحه (IRAC) -----------
def _format_brief_irac(case: Dict[str, object], citations: List[Dict[str, object]] | None) -> str:
    parties = case.get('parties') or {}
    title = case.get('title') or 'پیش‌نویس لایحه'
    facts = case.get('facts') or ''
    claims = case.get('claims') or ''
    court = case.get('court') or ''
    requests = case.get('requests') or ''
    # IRAC: Issue, Rule, Analysis, Conclusion
    lines = []
    lines.append(title)
    if court:
        lines.append(f"مرجع رسیدگی: {court}")
    if parties:
        lines.append(f"اصحاب دعوا: {parties}")
    lines.append("\n[مسئله]")
    lines.append(str(case.get('issue') or claims))
    lines.append("\n[قواعد (مواد قانونی و آراء مرتبط)]")
    if citations:
        for c in citations:
            lines.append(f"- {c.get('snippet')}")
    else:
        lines.append("- در حال حاضر استنادی ثبت نشده است.")
    lines.append("\n[تحلیل]")
    lines.append(str(case.get('analysis') or facts))
    lines.append("\n[نتیجه]")
    lines.append(str(case.get('conclusion') or requests))
    lines.append("\n[پیوست‌ها]")
    att = case.get('attachments') or []
    if att:
        for i, a in enumerate(att, 1):
            lines.append(f"ضمیمه {i}: {a}")
    return "\n".join([str(x) for x in lines])

@app.route('/draft', methods=['POST'])
def create_draft():
    data = request.get_json(silent=True) or {}
    question = data.get('prompt') or data.get('question') or 'تهیه پیش‌نویس لایحه'
    case = data.get('case') or {}
    # بازیابی استناد بر اساس عنوان/ادعا/وقایع
    needle = " ".join([str(case.get('title') or ''), str(case.get('claims') or ''), str(case.get('facts') or '')]).strip()
    citations = []
    try:
        cid = request.cookies.get('client_id')
        if cid and needle:
            _ensure_session_index(cid)
            retrieved = _retrieve_paragraphs(needle, cid, top_k=6)
            total_score = sum(int(r.get('score') or 0) for r in retrieved)
            strong = [r for r in retrieved if (r.get('score') or 0) >= 2]
            if retrieved and (len(strong) >= 1 or total_score >= 3):
                citations = [{'score': int(r['score']), 'snippet': (r['snippet'][:500] + '...') if len(r['snippet']) > 500 else r['snippet']} for r in retrieved]
                context = "\n\n---\n\n".join([c['snippet'] for c in citations])
            else:
                context = ''
        else:
            context = ''
    except Exception:
        context = ''
    # تولید متن لایحه: اگر مدل هست از generate_answer استفاده می‌کنیم، وگرنه فرمت IRAC محلی
    if context:
        prompt = f"بر اساس استنادات زیر لایحه‌ای رسمی و کوتاه با ساختار مسئله/قواعد/تحلیل/نتیجه بنویس:\n\n{context}\n\n[موضوع پرونده]\n{needle}\n\n[پاسخ]"
        draft_text = generate_answer(prompt, context)
    else:
        draft_text = _format_brief_irac(case, citations)
    return jsonify({ 'draft': draft_text, 'citations': citations })

@app.route('/draft/docx', methods=['POST'])
def create_draft_docx():
    try:
        from docx import Document
    except Exception:
        return jsonify({'error': 'کتابخانه python-docx نصب نیست. نصب: pip install python-docx'}), 400
    data = request.get_json(silent=True) or {}
    draft = data.get('draft') or ''
    if not draft:
        return jsonify({'error': 'متن لایحه ارسال نشده است'}), 400
    doc = Document()
    for line in str(draft).split('\n'):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    resp = make_response(bio.read())
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    resp.headers['Content-Disposition'] = 'attachment; filename="draft.docx"'
    return resp

def load_llama():
    global llama_instance, llama_error_message
    if llama_instance is not None:
        return llama_instance
    with _llama_lock:
        if llama_instance is not None:
            return llama_instance
        try:
            from llama_cpp import Llama
        except Exception as exc:
            llama_error_message = "کتابخانه llama-cpp-python نصب نیست. نصب: pip install llama-cpp-python"
            return None

        model_path = os.getenv('LLAMA_MODEL')
        if not model_path or not os.path.exists(model_path):
            llama_error_message = "مسیر مدل LLAMA_MODEL تنظیم/موجود نیست. یک فایل .gguf معتبر مشخص کنید."
            return None

        def _as_int(env_name, default_val):
            try:
                return int(os.getenv(env_name, str(default_val)))
            except ValueError:
                return default_val

        n_ctx = _as_int('LLAMA_CTX', 4096)
        n_threads = _as_int('LLAMA_THREADS', max(1, (os.cpu_count() or 4)))
        n_gpu_layers = _as_int('LLAMA_N_GPU_LAYERS', 0)

        try:
            kwargs = {
                'model_path': model_path,
                'n_ctx': n_ctx,
                'n_threads': n_threads
            }
            if n_gpu_layers > 0:
                kwargs['n_gpu_layers'] = n_gpu_layers
            llama_instance = Llama(**kwargs)
            llama_error_message = None
            return llama_instance
        except Exception as exc:
            llama_error_message = f"خطا در بارگذاری مدل: {exc}"
            return None

def generate_answer_with_llama(question, context):
    model = load_llama()
    if model is None:
        fallback = (
            "مدل محلی پیکربندی/نصب نشده است. "
            + (f"جزئیات: {llama_error_message}" if llama_error_message else "")
            + "\n\n" 
            + f"🔍 تحلیل اولیه بر اساس متن موجود:\n\n{question}\n\n👉 پاسخ نمونه: بر اساس متن داده‌شده، تشخیص دقیق نیاز به بررسی بیشتر دارد."
        )
        return fallback

    def _as_float(env_name, default_val):
        try:
            return float(os.getenv(env_name, str(default_val)))
        except ValueError:
            return default_val

    def _as_int(env_name, default_val):
        try:
            return int(os.getenv(env_name, str(default_val)))
        except ValueError:
            return default_val

    fast = _is_fast_mode()
    temperature = _as_float('LLAMA_TEMP', 0.15 if fast else 0.2)
    max_tokens = _as_int('LLAMA_MAX_TOKENS', 256 if fast else 512)
    top_p = _as_float('LLAMA_TOP_P', 0.9 if fast else 0.95)

    prompt = (
        "شما یک دستیار حقوقی فارسی هستید. با تکیه بر متن زمینه زیر، پاسخ دقیق و کوتاه بده.\n\n"
        + "[متن زمینه]" + "\n" + context + "\n\n"
        + "[سؤال]" + "\n" + question + "\n\n"
        + "[پاسخ]"
    )

    try:
        resp = model.create_completion(
            prompt=prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
        )
        text = resp.get('choices', [{}])[0].get('text', '').strip()
        return text or "پاسخی از مدل دریافت نشد."
    except Exception as exc:
        return f"خطا در تولید پاسخ از مدل: {exc}"

def generate_answer_with_ollama(question, context):
    import json
    import urllib.request
    import urllib.error

    host = os.getenv('OLLAMA_HOST', 'http://127.0.0.1:11434')
    model = os.getenv('OLLAMA_MODEL', 'gemma3:4b')
    url = host.rstrip('/') + '/api/generate'

    prompt = (
        "شما یک دستیار حقوقی فارسی هستید. با تکیه بر متن زمینه زیر، پاسخ دقیق و کوتاه بده.\n\n"
        + "[متن زمینه]" + "\n" + context + "\n\n"
        + "[سؤال]" + "\n" + question + "\n\n"
        + "[پاسخ]"
    )

    fast = _is_fast_mode()
    payload = {
        'model': model,
        'prompt': prompt,
        'stream': False,
        'options': {
            'temperature': float(os.getenv('OLLAMA_TEMP', '0.15' if fast else '0.2')),
            'top_p': float(os.getenv('OLLAMA_TOP_P', '0.9' if fast else '0.95')),
            'num_predict': int(os.getenv('OLLAMA_MAX_TOKENS', '256' if fast else '512')),
        }
    }

    def list_installed_models():
        try:
            tags_url = host.rstrip('/') + '/api/tags'
            with urllib.request.urlopen(tags_url, timeout=5 if _is_fast_mode() else 10) as resp:
                tags = json.loads(resp.read().decode('utf-8')) or {}
                models = tags.get('models') or []
                names = []
                for m in models:
                    name = m.get('name') or m.get('model')
                    if name:
                        names.append(name)
                return names
        except Exception:
            return []

    def is_reachable():
        try:
            tags_url = host.rstrip('/') + '/api/tags'
            with urllib.request.urlopen(tags_url, timeout=3) as resp:
                _ = resp.read()
                return True
        except Exception:
            return False

    def do_chat(use_model):
        chat_url = host.rstrip('/') + '/api/chat'
        messages = [
            { 'role': 'system', 'content': 'شما یک دستیار حقوقی فارسی هستید. با تکیه بر متن زمینه زیر، پاسخ دقیق و کوتاه بده.' },
            { 'role': 'user', 'content': f"[متن زمینه]\n{context}\n\n[سؤال]\n{question}" }
        ]
        payload_chat = {
            'model': use_model,
            'messages': messages,
            'stream': False,
            'options': {
                'temperature': float(os.getenv('OLLAMA_TEMP', '0.2')),
                'top_p': float(os.getenv('OLLAMA_TOP_P', '0.95')),
                'num_predict': int(os.getenv('OLLAMA_MAX_TOKENS', '512')),
            }
        }
        data_chat = json.dumps(payload_chat).encode('utf-8')
        req_chat = urllib.request.Request(chat_url, data=data_chat, headers={'Content-Type': 'application/json'})
        try:
            try:
                timeout_s = float(os.getenv('OLLAMA_TIMEOUT', '45' if fast else '120'))
            except ValueError:
                timeout_s = 45.0 if fast else 120.0
            with urllib.request.urlopen(req_chat, timeout=timeout_s) as resp:
                res = json.loads(resp.read().decode('utf-8'))
                # ساختار chat: res['message']['content']
                msg = (res.get('message') or {}).get('content', '')
                return msg.strip() or "پاسخی از مدل دریافت نشد."
        except Exception as exc:
            return f"خطا در ارتباط با Ollama (chat): {exc}"

    def do_generate(use_model, tried_fallback=False):
        payload_local = dict(payload)
        payload_local['model'] = use_model
        data_local = json.dumps(payload_local).encode('utf-8')
        req_local = urllib.request.Request(url, data=data_local, headers={'Content-Type': 'application/json'})
        try:
            try:
                timeout_s = float(os.getenv('OLLAMA_TIMEOUT', '45' if fast else '120'))
            except ValueError:
                timeout_s = 45.0 if fast else 120.0
            with urllib.request.urlopen(req_local, timeout=timeout_s) as resp:
                res = json.loads(resp.read().decode('utf-8'))
                text = (res.get('response') or '').strip()
                return text or "پاسخی از مدل دریافت نشد."
        except urllib.error.HTTPError as http_exc:
            # اگر مدل پیدا نشد (404)، روی یک مدل نصب‌شده سوییچ کن
            if http_exc.code == 404:
                # ابتدا تلاش با /api/chat برای برخی خانواده‌ها (مثل qwen3/deepseek)
                chat_ans = do_chat(use_model)
                if chat_ans and not chat_ans.startswith('خطا در ارتباط با Ollama'):
                    return chat_ans
                installed = list_installed_models()
                fallback = os.getenv('OLLAMA_FALLBACK_MODEL')
                pick = fallback or ('gemma3:4b' if 'gemma3:4b' in installed else ('llama3.2' if 'llama3.2' in installed else (installed[0] if installed else None)))
                if pick and pick != use_model:
                    try:
                        return do_generate(pick, tried_fallback=True)
                    except Exception:
                        pass
                return f"مدل '{use_model}' در Ollama پیدا نشد. یک مدل نصب کنید (مثلاً: ollama pull {use_model}) یا متغیر OLLAMA_MODEL را روی یکی از نصب‌شده‌ها تنظیم کنید: {installed}"
            return f"خطای HTTP از Ollama: {http_exc}"
        except Exception as exc:
            # برخی محیط‌ها ممکن است به جای HTTPError، Exception عمومی با متن 404 بدهند
            msg = str(exc)
            if (('404' in msg) or ('Not Found' in msg)) and not tried_fallback:
                chat_ans = do_chat(use_model)
                if chat_ans and not chat_ans.startswith('خطا در ارتباط با Ollama'):
                    return chat_ans
                installed = list_installed_models()
                fallback = os.getenv('OLLAMA_FALLBACK_MODEL')
                pick = fallback or ('gemma3:4b' if 'gemma3:4b' in installed else ('llama3.2' if 'llama3.2' in installed else (installed[0] if installed else None)))
                if pick and pick != use_model:
                    try:
                        return do_generate(pick, tried_fallback=True)
                    except Exception:
                        pass
            return f"خطا در ارتباط با Ollama: {exc}"

    # اگر سرویس در دسترس نیست، سریعاً خارج شو (بدون وابستگی به نصب مدل)
    if not is_reachable():
        return 'خطا در ارتباط با Ollama: سرویس در دسترس نیست'

    # ابتدا chat را امتحان می‌کنیم (برای خانواده‌هایی مثل gemma/deepseek/qwen)
    chat_first = do_chat(model)
    if chat_first and not chat_first.startswith('خطا در ارتباط با Ollama'):
        return chat_first
    # در صورت شکست chat، generate را امتحان می‌کنیم با مکانیزم fallback
    return do_generate(model)

def _should_use_deepseek() -> bool:
    key = os.getenv('DEEPSEEK_API_KEY')
    if key and len(key.strip()) > 10:
        return True
    return _str_to_bool(os.getenv('USE_DEEPSEEK', '0'), default=False)

def generate_answer_with_deepseek(question: str, context: str) -> str:
    try:
        import requests  # type: ignore
    except Exception:
        return 'خطا در DeepSeek: کتابخانه requests در دسترس نیست'

    api_key = os.getenv('DEEPSEEK_API_KEY', '').strip()
    model = os.getenv('DEEPSEEK_MODEL', 'deepseek-chat').strip() or 'deepseek-chat'
    if not api_key:
        return 'خطا در DeepSeek: کلید API تنظیم نشده است'

    fast = _is_fast_mode()
    try:
        temperature = float(os.getenv('DEEPSEEK_TEMP', '0.15' if fast else '0.2'))
    except Exception:
        temperature = 0.15 if fast else 0.2
    try:
        top_p = float(os.getenv('DEEPSEEK_TOP_P', '0.9' if fast else '0.95'))
    except Exception:
        top_p = 0.9 if fast else 0.95
    try:
        max_tokens = int(os.getenv('DEEPSEEK_MAX_TOKENS', '256' if fast else '512'))
    except Exception:
        max_tokens = 256 if fast else 512

    url = 'https://api.deepseek.com/chat/completions'
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    payload = {
        'model': model,
        'messages': [
            { 'role': 'system', 'content': 'شما یک دستیار حقوقی فارسی هستید. با تکیه بر متن زمینه زیر، پاسخ دقیق و کوتاه بده.' },
            { 'role': 'user', 'content': f"[متن زمینه]\n{context}\n\n[سؤال]\n{question}" },
        ],
        'temperature': temperature,
        'top_p': top_p,
        'max_tokens': max_tokens,
        'stream': False,
    }
    try:
        timeout_s = float(os.getenv('DEEPSEEK_TIMEOUT', '45' if fast else '120'))
    except Exception:
        timeout_s = 45.0 if fast else 120.0

    ts0 = time.perf_counter()
    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=timeout_s)
        took = int((time.perf_counter() - ts0) * 1000)
        MODEL_STATUS.update({'engine':'deepseek','latency_ms':took,'ts':int(time.time())})
        if resp.status_code >= 400:
            MODEL_STATUS.update({'error': f'HTTP {resp.status_code}'})
            return f"خطا در DeepSeek: HTTP {resp.status_code} - {resp.text[:200]}"
        data = resp.json()
        choices = data.get('choices') or []
        if not choices:
            MODEL_STATUS.update({'error':'no_choices'})
            return 'پاسخی از DeepSeek دریافت نشد.'
        content = ((choices[0] or {}).get('message') or {}).get('content', '')
        MODEL_STATUS.update({'error': ''})
        return content.strip() or 'پاسخی از DeepSeek دریافت نشد.'
    except Exception as exc:
        MODEL_STATUS.update({'engine':'deepseek','error': str(exc), 'ts': int(time.time())})
        return f"خطا در ارتباط با DeepSeek: {exc}"

def generate_answer(question, context):
    # ابتدا پاسخ‌های بازبینی‌شدهٔ انسانی را چک می‌کنیم
    try:
        key = _hash_key_for_question(question)
        curated = CURATED_ANSWERS.get(key)
        if curated and curated.get('answer'):
            return str(curated['answer'])
    except Exception:
        pass
    # حالت بدون مدل (اختیاری): فقط پاسخ محلی تولید شود
    if _str_to_bool(os.getenv('NO_MODELS', '0'), default=False):
        return generate_answer_locally(question, context)

    # تلاش با DeepSeek (در صورت پیکربندی)
    deepseek_ans = None
    if _should_use_deepseek():
        deepseek_ans = generate_answer_with_deepseek(question, context)
        if isinstance(deepseek_ans, str) and not deepseek_ans.startswith('خطا در') and not deepseek_ans.startswith('پاسخی از DeepSeek'):
            return deepseek_ans

    # تلاش با llama-cpp (در صورت فعال بودن)
    llama_ans = None
    if _should_use_llama():
        llama_ans = generate_answer_with_llama(question, context)
        if isinstance(llama_ans, str) and not llama_ans.startswith('مدل محلی پیکربندی') and not llama_ans.startswith('خطا در'):
            return llama_ans

    # تلاش با Ollama
    ollama_ans = None
    if _should_use_ollama():
        ollama_ans = generate_answer_with_ollama(question, context)
        if isinstance(ollama_ans, str) and not ollama_ans.startswith('خطا در ارتباط با Ollama') and not ollama_ans.startswith('خطای HTTP') and not ollama_ans.startswith('خطا در'):
            return ollama_ans

    # اگر هر دو مسیر شکست خورد، پاسخ محلی بده و پیام‌های خطا را ضمیمه کن
    local = generate_answer_locally(question, context)
    debug = []
    if isinstance(deepseek_ans, str) and (deepseek_ans.startswith('خطا در') or deepseek_ans.startswith('پاسخی از DeepSeek')):
        debug.append(deepseek_ans)
    if isinstance(llama_ans, str) and (llama_ans.startswith('مدل محلی پیکربندی') or llama_ans.startswith('خطا در')):
        debug.append(llama_ans)
    if isinstance(ollama_ans, str) and (ollama_ans.startswith('خطا در ارتباط با Ollama') or ollama_ans.startswith('خطای HTTP') or ollama_ans.startswith('خطا در')):
        debug.append(ollama_ans)
    return local + ("\n\n" + "\n".join(debug) if debug else '')

# ----------- صفحه وب ساده -----------
@app.route('/', methods=['GET'])
def index():
    # خواندن خودکار فقط اگر هنوز چیزی برای این سشن پردازش نشده باشد
    client_id = request.cookies.get('client_id') or uuid.uuid4().hex
    with _context_lock:
        existing = SESSION_CONTEXTS.get(client_id) or []
        already_processed = SESSION_PROCESSED.get(client_id) or set()
    saved_texts: List[str] = []
    previews: List[str] = []
    total_files = 0
    if not existing and not already_processed:
        saved_texts, previews, total_files = ingest_pdfs_from_dir(DEFAULT_PDF_DIR, recursive=True, client_id=client_id)
    if saved_texts:
        _append_session_context(client_id, saved_texts)
    combined_preview = '\n\n---\n\n'.join(previews) if previews else ''

    html = render_template('index.html', combined_preview=combined_preview, total_files=total_files)

    resp = make_response(html)
    # ست کردن کوکی سشن در صورت نبود
    if not request.cookies.get('client_id'):
        resp.set_cookie('client_id', client_id, max_age=30*24*3600, httponly=False, samesite='Lax')
    return resp

# آپلود حذف شده است

# ----------- خواندن PDFها از پوشهٔ روی سرور -----------
@app.route('/ingest-dir', methods=['POST'])
def ingest_directory():
    data = request.get_json(silent=True) or {}
    dir_path = data.get('dir')
    recursive = bool(data.get('recursive', True))
    # اندازهٔ بچ (برای بارگذاری تکه‌تکه)
    try:
        req_max = int(data.get('max_files', 0))
    except Exception:
        req_max = 0
    # اگر dir ارسال نشده باشد، از پوشهٔ پیش‌فرض استفاده می‌کنیم
    if not dir_path:
        dir_path = DEFAULT_PDF_DIR
        recursive = True
    if not dir_path or not isinstance(dir_path, str):
        return jsonify({'error': 'مسیر پوشه ارسال نشده است'}), 400
    # نرمال‌سازی مسیر (پشتیبانی از مسیرهای فارسی/ویندوز)
    dir_path = os.path.expanduser(dir_path)
    if not os.path.isabs(dir_path):
        dir_path = os.path.abspath(os.path.join(BASE_DIR, dir_path))
    if not os.path.isdir(dir_path):
        return jsonify({'error': 'پوشه یافت نشد'}), 400

    client_id = _get_or_create_client_id()
    saved_texts: List[str] = []
    previews: List[str] = []
    total_files = 0
    # حداکثر هر فراخوانی
    max_files = 200
    if req_max and req_max > 0:
        max_files = max(1, min(req_max, 500))

    def iter_pdfs(path):
        if recursive:
            for root, _dirs, files in os.walk(path):
                for name in files:
                    if name.lower().endswith('.pdf'):
                        yield os.path.join(root, name)
        else:
            for name in os.listdir(path):
                fp = os.path.join(path, name)
                if os.path.isfile(fp) and name.lower().endswith('.pdf'):
                    yield fp

    fast = bool(data.get('fast', False))
    for fp in iter_pdfs(dir_path):
        try:
            extracted_text = extract_text_fast(fp) if fast else extract_text(fp)
        except Exception:
            extracted_text = ''
        if extracted_text:
            saved_texts.append(extracted_text)
            pv = (extracted_text[:500] + '...') if len(extracted_text) > 500 else extracted_text
            previews.append(pv)
            total_files += 1
            if total_files >= max_files:
                break

    if saved_texts:
        _append_session_context(client_id, saved_texts)
        # پس از به‌روزرسانی متن‌ها، ایندکس این سشن را نیز تازه‌سازی کنیم
        try:
            _ensure_session_index(client_id)
        except Exception:
            pass

    combined_preview = '\n\n---\n\n'.join(previews)
    resp = jsonify({
        'message': 'فایل‌ها از پوشه پردازش شدند.',
        'dir': dir_path,
        'total_files': total_files,
        'previews': previews,
        'combined_preview': combined_preview
    })
    response = make_response(resp)
    response.set_cookie('client_id', client_id, max_age=30*24*3600, httponly=False, samesite='Lax')
    return response

@app.route('/ingest-prewarm', methods=['POST'])
def ingest_prewarm():
    """شروع بارگذاری تکه‌تکه در پس‌زمینه برای شروع سریع."""
    data = request.get_json(silent=True) or {}
    dir_path = data.get('dir') or DEFAULT_PDF_DIR
    recursive = bool(data.get('recursive', True))
    try:
        batch_size = int(data.get('max_files', 20))
    except Exception:
        batch_size = 20
    fast = _str_to_bool(str(data.get('fast', '1')), default=True)

    client_id = _get_or_create_client_id()

    # محاسبهٔ کل فایل‌های هدف (به‌جز پردازش‌شده‌های قبلی این سشن)
    def _count_total_targets() -> int:
        try:
            proc: Set[str] = set()
            with _context_lock:
                proc = set(SESSION_PROCESSED.get(client_id) or set())
            cnt = 0
            if recursive:
                for root, _dirs, files in os.walk(dir_path):
                    for name in files:
                        if name.lower().endswith('.pdf'):
                            fp = os.path.abspath(os.path.join(root, name))
                            if fp not in proc:
                                cnt += 1
            else:
                for name in os.listdir(dir_path):
                    fp = os.path.abspath(os.path.join(dir_path, name))
                    if os.path.isfile(fp) and name.lower().endswith('.pdf') and fp not in proc:
                        cnt += 1
            return cnt
        except Exception:
            return 0

    total_targets = _count_total_targets()
    with _context_lock:
        PREWARM_PROGRESS[client_id] = {
            'running': True,
            'done': 0,
            'total': total_targets,
            'dir': dir_path,
            'fast': fast,
        }

    def _worker():
        try:
            while True:
                _texts, _prev, count = ingest_pdfs_from_dir(dir_path, recursive=recursive, client_id=client_id, fast=fast, max_files=batch_size)
                if count <= 0:
                    with _context_lock:
                        st = PREWARM_PROGRESS.get(client_id) or {}
                        st['running'] = False
                        PREWARM_PROGRESS[client_id] = st
                    break
                with _context_lock:
                    st = PREWARM_PROGRESS.get(client_id) or {}
                    st_done = int(st.get('done') or 0) + int(count)
                    st['done'] = st_done
                    PREWARM_PROGRESS[client_id] = st
        except Exception:
            pass

    t = threading.Thread(target=_worker, daemon=True)
    t.start()
    resp = jsonify({'message': 'پیش‌بارگذاری آغاز شد', 'client_id': client_id})
    response = make_response(resp)
    response.set_cookie('client_id', client_id, max_age=30*24*3600, httponly=False, samesite='Lax')
    return response

@app.route('/ingest-progress', methods=['GET'])
def ingest_progress():
    client_id = request.cookies.get('client_id')
    with _context_lock:
        st = PREWARM_PROGRESS.get(client_id or '')
    if not st:
        return jsonify({'running': False, 'done': 0, 'total': 0, 'percent': 0})
    done = int(st.get('done') or 0)
    total = int(st.get('total') or 0)
    percent = int((done * 100) / total) if total > 0 else (100 if not st.get('running') else 0)
    return jsonify({
        'running': bool(st.get('running')),
        'done': done,
        'total': total,
        'percent': percent,
        'dir': st.get('dir'),
        'fast': st.get('fast'),
    })

# مدیریت خطای حجم زیاد (413) به صورت JSON
@app.errorhandler(413)
def request_entity_too_large(_):
    return jsonify({'error': 'حجم فایل بیش از حد مجاز است (حداکثر ۱ گیگابایت).'}), 413

# پاسخ JSON برای همهٔ خطاهای HTTP (جلوگیری از پاسخ HTML که باعث Failed to fetch در fetch می‌شود)
@app.errorhandler(HTTPException)
def handle_http_exception(e: HTTPException):
    response = e.get_response()
    return jsonify({'error': e.description or 'خطای سرور'}), e.code

# سلامت
@app.get('/healthz')
def healthz():
    return 'ok', 200

# ----------- مسیر متن کامل سشن -----------
@app.route('/context', methods=['GET'])
def get_full_context():
    client_id = request.cookies.get('client_id')
    texts: List[str] = []
    if client_id:
        with _context_lock:
            texts = SESSION_CONTEXTS.get(client_id) or []
    # اگر سشن خالی بود، به‌صورت خودکار از پوشهٔ پیش‌فرض بخوان
    if not texts:
        cid = client_id or uuid.uuid4().hex
        logger.info("Context empty; auto-ingest from default: %s", DEFAULT_PDF_DIR)
        saved_texts, _previews, _total = ingest_pdfs_from_dir(DEFAULT_PDF_DIR, recursive=True, client_id=cid)
        if saved_texts:
            _append_session_context(cid, saved_texts)
            client_id = cid
            with _context_lock:
                texts = SESSION_CONTEXTS.get(client_id) or []
    combined = "\n\n---\n\n".join(texts)
    if request.args.get('download'):
        # دانلود به صورت فایل متنی
        resp = make_response(combined)
        resp.headers['Content-Type'] = 'text/plain; charset=utf-8'
        resp.headers['Content-Disposition'] = 'attachment; filename="context.txt"'
        # اطمینان از ست‌شدن کوکی سشن
        if client_id and not request.cookies.get('client_id'):
            resp.set_cookie('client_id', client_id, max_age=30*24*3600, httponly=False, samesite='Lax')
        return resp
    data = {'combined': combined, 'count': len(texts)}
    resp_json = make_response(jsonify(data))
    if client_id and not request.cookies.get('client_id'):
        resp_json.set_cookie('client_id', client_id, max_age=30*24*3600, httponly=False, samesite='Lax')
    return resp_json

# ----------- مسیر پرسیدن سؤال -----------
@app.route('/ask', methods=['POST'])
@limiter.limit(os.getenv('RATE_LIMIT_ASK', '60 per minute'))
def ask_question():
    data = request.get_json()
    question = data.get('question')
    context = data.get('context')

    if not question:
        return jsonify({'error': 'سؤال ارسال نشده است'}), 400

    # اگر Context خالی بود، از حافظهٔ سشن (بر پایهٔ کوکی client_id) استفاده می‌کنیم
    if not context or not context.strip():
        client_id = request.cookies.get('client_id')
        if client_id:
            with _context_lock:
                texts = SESSION_CONTEXTS.get(client_id) or []
            if texts:
                # محدودسازی زمینه برای سرعت بیشتر
                if _is_fast_mode():
                    context = _build_limited_context(texts, per_text_limit=1000, total_limit=8000)
                else:
                    context = _build_limited_context(texts, per_text_limit=2000, total_limit=15000)
        #Fallback: اگر هنوز خالی است، یک‌بار ingest از پوشهٔ پیش‌فرض انجام بده
        if not context or not context.strip():
            cid = client_id or uuid.uuid4().hex
            logger.info("/ask fallback ingest from default: %s", DEFAULT_PDF_DIR)
            saved_texts, _previews, _total = ingest_pdfs_from_dir(DEFAULT_PDF_DIR, recursive=True, client_id=cid)
            if saved_texts:
                _append_session_context(cid, saved_texts)
                with _context_lock:
                    texts = SESSION_CONTEXTS.get(cid) or []
                if texts:
                    if _is_fast_mode():
                        context = _build_limited_context(texts, per_text_limit=1000, total_limit=8000)
                    else:
                        context = _build_limited_context(texts, per_text_limit=2000, total_limit=15000)
                # ست‌کردن کوکی سشن در پاسخ
                resp = jsonify({'error': None})
                resp.set_cookie('client_id', cid, max_age=30*24*3600, httponly=False, samesite='Lax')
                # ادامهٔ پردازش پاسخ با context پر شده؛ بدون برگرداندن این resp موقت
    if not context or not context.strip():
        # اگر متن زمینه مهیا نشد، با متن خالی ادامه می‌دهیم تا پاسخ حداقلی برگردد
        context = ""

    # RAG: ساخت زمینه بر اساس بازیابی پاراگراف‌ها، با دروازهٔ استناد
    citations = []
    try:
        client_id_eff = request.cookies.get('client_id')
        if client_id_eff:
            _ensure_session_index(client_id_eff)
            retrieved = _retrieve_paragraphs(question, client_id_eff, top_k=5)
            # گیت ساده: حداقل 2 پاراگراف با امتیاز >= 2 یا مجموع امتیازها >= 3
            total_score = sum(int(r.get('score') or 0) for r in retrieved)
            strong = [r for r in retrieved if (r.get('score') or 0) >= 2]
            if retrieved and (len(strong) >= 1 or total_score >= 3):
                citations = [{'score': int(r['score']), 'snippet': (r['snippet'][:400] + '...') if len(r['snippet']) > 400 else r['snippet']} for r in retrieved]
                rag_context = "\n\n---\n\n".join([c['snippet'] for c in citations])
                # استفاده از زمینهٔ بازیابی‌شده به‌جای زمینهٔ خام
                context = rag_context
    except Exception:
        pass

    answer = generate_answer(question, context)
    return jsonify({'answer': answer, 'citations': citations})

# ----------- استریم پاسخ (SSE) -----------
@app.route('/ask-stream', methods=['POST'])
def ask_question_stream():
    data = request.get_json()
    question = (data or {}).get('question')
    context = (data or {}).get('context') or ''
    if not question:
        def _bad():
            yield 'data: ' + '{"error":"سؤال ارسال نشده است"}' + '\n\n'
        return Response(stream_with_context(_bad()), mimetype='text/event-stream')

    def _gen():
        import json as _json
        yield 'data: ' + _json.dumps({'type': 'status', 'message': 'retrieving'}) + '\n\n'
        citations = []
        try:
            cid = request.cookies.get('client_id')
            if cid:
                _ensure_session_index(cid)
                retrieved = _retrieve_paragraphs(question, cid, top_k=5)
                total_score = sum(int(r.get('score') or 0) for r in retrieved)
                strong = [r for r in retrieved if (r.get('score') or 0) >= 2]
                if retrieved and (len(strong) >= 1 or total_score >= 3):
                    citations = [{'score': int(r['score']), 'snippet': (r['snippet'][:400] + '...') if len(r['snippet']) > 400 else r['snippet']} for r in retrieved]
                    context_local = "\n\n---\n\n".join([c['snippet'] for c in citations])
                else:
                    context_local = context
            else:
                context_local = context
        except Exception:
            context_local = context
        yield 'data: ' + _json.dumps({'type': 'citations', 'items': citations}) + '\n\n'
        ans = generate_answer(question, context_local)
        yield 'data: ' + _json.dumps({'type': 'answer', 'text': ans}) + '\n\n'
    return Response(stream_with_context(_gen()), mimetype='text/event-stream')

# ----------- پنل ادمین -----------
@app.route('/admin', methods=['GET'])
def admin_panel():
    return render_template('admin.html')

@app.route('/admin/stats', methods=['GET'])
def admin_stats():
    # آمار سادهٔ سشن‌ها و فایل‌های پردازش‌شده
    with _context_lock:
        session_count = len(SESSION_CONTEXTS)
        total_texts = TOTAL_TEXTS_COUNT
        processed_files = PROCESSED_FILES_COUNT
    # وضعیت Ollama
    ollama_status = 'disabled'
    try:
        if _should_use_ollama():
            import json as _json
            import urllib.request as _u
            host = os.getenv('OLLAMA_HOST', 'http://127.0.0.1:11434')
            with _u.urlopen(host.rstrip('/') + '/api/tags', timeout=3) as r:
                _ = _json.loads(r.read().decode('utf-8'))
                ollama_status = 'online'
        else:
            ollama_status = 'disabled'
    except Exception:
        ollama_status = 'offline'

    # وضعیت LLaMA (بدون بارگذاری مدل)
    llama_status = 'disabled'
    use_llama = _should_use_llama()
    llama_model = os.getenv('LLAMA_MODEL', '')
    if use_llama:
        if llama_model and os.path.exists(llama_model):
            llama_status = 'configured'
        else:
            llama_status = 'missing_model'

    # DeepSeek وضعیت
    deepseek_key = os.getenv('DEEPSEEK_API_KEY', '').strip()
    deepseek_status = 'configured' if deepseek_key else 'disabled'
    deepseek_model = os.getenv('DEEPSEEK_MODEL', 'deepseek-chat')

    # انتخاب موتور فعال: DeepSeek > LLaMA > Ollama > local
    engine = 'local'
    if deepseek_status == 'configured':
        engine = 'deepseek'
    elif use_llama and llama_status == 'configured':
        engine = 'llama'
    elif _should_use_ollama() and ollama_status == 'online':
        engine = 'ollama'

    # وضعیت برداری
    vector_ready = GLOBAL_VECTOR.get('index') is not None or (os.path.exists(os.path.join(VECTOR_INDEX_DIR, 'index.faiss')))

    data = {
        'sessions': session_count,
        'texts_in_memory': total_texts,
        'processed_files': processed_files,
        'default_pdf_dir': DEFAULT_PDF_DIR,
        'fast_mode': _is_fast_mode(),
        'ollama': ollama_status,
        'ollama_host': os.getenv('OLLAMA_HOST', 'http://127.0.0.1:11434'),
        'ollama_model': os.getenv('OLLAMA_MODEL', 'gemma3:4b'),
        'llama': llama_status,
        'llama_model': llama_model,
        'deepseek': deepseek_status,
        'deepseek_model': deepseek_model,
        'engine': engine,
        'model_status': MODEL_STATUS,
        'vector_ready': bool(vector_ready),
        'vector_model': GLOBAL_VECTOR.get('model_name') or os.getenv('EMBED_MODEL', ''),
    }
    return jsonify(data)

@app.route('/admin/config', methods=['GET', 'POST'])
def admin_config():
    if request.method == 'GET':
        cfg = _read_config_file()
        # ماسک کردن کلیدها
        masked = dict(cfg)
        if masked.get('DEEPSEEK_API_KEY'):
            masked['DEEPSEEK_API_KEY'] = _mask_secret(str(masked['DEEPSEEK_API_KEY']))
        return jsonify(masked)
    # POST
    data = request.get_json(silent=True) or {}
    cfg = _read_config_file()
    # فقط کلیدهای مجاز
    allowed = {
        'DEEPSEEK_API_KEY', 'DEEPSEEK_MODEL',
        'USE_OLLAMA', 'USE_LLAMA', 'USE_DEEPSEEK'
    }
    for k, v in (data.items() if isinstance(data, dict) else []):
        if k in allowed:
            cfg[k] = v
    ok = _write_config_file(cfg)
    if ok:
        _load_persisted_config_into_env()
        return jsonify({'ok': True})
    return jsonify({'ok': False}), 500

@app.route('/admin/clear', methods=['POST'])
def admin_clear():
    # پاک‌سازی حافظهٔ موقت سشن‌ها
    with _context_lock:
        SESSION_CONTEXTS.clear()
        SESSION_PROCESSED.clear()
    return jsonify({'message': 'حافظهٔ سشن‌ها پاک شد.'})

@app.route('/admin/reindex', methods=['POST'])
def admin_reindex():
    # بازخوانی پوشهٔ پیش‌فرض و ذخیرهٔ متن‌ها برای یک سشن جدید موقتی
    client_id = uuid.uuid4().hex
    saved_texts, previews, total_files = ingest_pdfs_from_dir(DEFAULT_PDF_DIR, recursive=True, client_id=client_id)
    if saved_texts:
        _append_session_context(client_id, saved_texts)
    return jsonify({
        'message': 'بازخوانی انجام شد.',
        'client_id': client_id,
        'total_files': total_files,
        'previews_count': len(previews)
    })

@app.route('/admin/build-vector', methods=['POST'])
def admin_build_vector():
    # ساخت ایندکس برداری از قوانین/متون موجود
    try:
        res = build_global_vector_index()
        return jsonify(res)
    except Exception as exc:
        return jsonify({'built': False, 'reason': str(exc)}), 500

@app.route('/admin/web-ingest', methods=['POST'])
def admin_web_ingest():
    """دریافت منابع عمومی از وب (URLها)، استخراج متن و ذخیره در پوشهٔ sources."""
    data = request.get_json(silent=True) or {}
    urls = data.get('urls') or []
    add_to_vector = bool(data.get('add_to_vector', False))
    timeout_sec = float(data.get('timeout', 15))
    if isinstance(urls, str):
        urls = [u.strip() for u in urls.split('\n') if u.strip()]
    if not isinstance(urls, list) or not urls:
        return jsonify({'error': 'فهرست URLs نامعتبر است'}), 400
    import urllib.request
    import hashlib
    saved_files = []
    texts = []
    for u in urls:
        try:
            req = urllib.request.Request(u, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
                raw = resp.read()
                try:
                    html = raw.decode('utf-8', errors='ignore')
                except Exception:
                    html = raw.decode(errors='ignore')
                text = extract_text_from_html(html)
                if text:
                    h = hashlib.sha1(u.encode('utf-8', errors='ignore')).hexdigest()[:12]
                    outp = os.path.join(SOURCES_DIR, f'{h}.txt')
                    with open(outp, 'w', encoding='utf-8') as f:
                        f.write(text)
                    saved_files.append(outp)
                    texts.append(text)
        except Exception:
            continue
    # در صورت درخواست، به ایندکس برداری موجود اضافه کن
    appended = 0
    if add_to_vector and texts and GLOBAL_VECTOR.get('index') is not None:
        model, _m = _get_embedder()
        if model is not None:
            import numpy as np  # type: ignore
            import faiss  # type: ignore
            paras_new = []
            for t in texts:
                paras_new.extend(_split_paragraphs(t))
            paras_new = [p for p in paras_new if p and len(p.strip()) >= 20]
            if paras_new:
                emb = model.encode(paras_new, convert_to_numpy=True, normalize_embeddings=True)
                try:
                    GLOBAL_VECTOR['index'].add(emb)
                    GLOBAL_VECTOR['paragraphs'].extend(paras_new)
                    appended = len(paras_new)
                except Exception:
                    appended = 0
    return jsonify({'saved': len(saved_files), 'files': saved_files, 'appended_to_vector': appended})

@app.route('/admin/curate', methods=['POST'])
def admin_curate_answer():
    """ثبت پاسخ بازبینی‌شده توسط انسان برای یک سؤال (HITL)."""
    data = request.get_json(silent=True) or {}
    question = (data.get('question') or '').strip()
    answer = (data.get('answer') or '').strip()
    if not question or not answer:
        return jsonify({'error': 'question/answer الزامی است'}), 400
    key = _hash_key_for_question(question)
    CURATED_ANSWERS[key] = { 'answer': answer, 'by': 'admin' }
    return jsonify({'ok': True, 'key': key})

@app.route('/admin/curate', methods=['GET'])
def admin_list_curated():
    out = []
    for k, v in CURATED_ANSWERS.items():
        out.append({ 'key': k, 'len': len(str(v.get('answer') or '')) })
    return jsonify({ 'count': len(out), 'items': out })

@app.route('/admin/set-deepseek', methods=['POST'])
def admin_set_deepseek():
    data = request.get_json(silent=True) or {}
    key = (data.get('api_key') or '').strip()
    model = (data.get('model') or '').strip() or 'deepseek-chat'
    if not key:
        return jsonify({'ok': False, 'error': 'api_key الزامی است'}), 400
    # ذخیره در فایل تنظیمات و بارگذاری مجدد در env
    cfg = _read_config_file()
    cfg['DEEPSEEK_API_KEY'] = key
    cfg['DEEPSEEK_MODEL'] = model
    cfg['USE_DEEPSEEK'] = '1'
    cfg['USE_OLLAMA'] = '0'
    cfg['USE_LLAMA'] = '0'
    if _write_config_file(cfg):
        _load_persisted_config_into_env()
        return jsonify({'ok': True, 'model': model})
    return jsonify({'ok': False}), 500

@app.route('/webhook', methods=['POST'])
@limiter.limit(os.getenv('RATE_LIMIT_WEBHOOK', '30 per minute'))
def webhook_receiver():
    # توکن وبهوک از هدر یا querystring خوانده می‌شود
    expected = os.getenv('WEBHOOK_TOKEN') or os.getenv('WEBHOOK_SECRET') or os.getenv('TOKEN')
    if expected:
        provided = (
            request.headers.get('X-Webhook-Token')
            or request.headers.get('X-Token')
            or request.args.get('token')
        )
        if not provided or provided != expected:
            return jsonify({'ok': False, 'error': 'unauthorized'}), 401

    # اختیاری: اعتبارسنجی امضا (HMAC-SHA256) روی بادی خام
    hsecret = os.getenv('WEBHOOK_HMAC_SECRET')
    if hsecret:
        try:
            raw = request.get_data(cache=True) or b''
            sig = hmac.new(hsecret.encode('utf-8'), raw, hashlib.sha256).hexdigest()
            provided_sig = request.headers.get('X-Signature') or request.headers.get('X-Hub-Signature-256')
            if not provided_sig or (sig.lower() != provided_sig.strip().lower().replace('sha256=','')):
                return jsonify({'ok': False, 'error': 'bad_signature'}), 401
        except Exception:
            return jsonify({'ok': False, 'error': 'sig_check_failed'}), 400

    payload = request.get_json(silent=True) or {}
    # ذخیرهٔ payload برای استفادهٔ آتی/اسکریپت
    try:
        import json as _json
        import time as _time
        save_dir = os.path.join(BASE_DIR, 'tmp')
        os.makedirs(save_dir, exist_ok=True)
        fname = f"webhook_{int(_time.time())}.json"
        fpath = os.path.join(save_dir, fname)
        with open(fpath, 'w', encoding='utf-8') as f:
            _json.dump(payload, f, ensure_ascii=False)
    except Exception:
        fpath = ''

    # اجرای اختیاری فایل .bat برای پردازش payload
    started = False
    bat_path = os.getenv('WEBHOOK_BAT')
    if bat_path and os.path.isfile(bat_path):
        try:
            import subprocess as _sub
            # اجرای غیربلاک‌کننده با ارسال مسیر فایل payload به عنوان آرگومان اول
            _sub.Popen(['cmd.exe', '/C', bat_path, fpath or ''], cwd=BASE_DIR)
            started = True
        except Exception:
            started = False

    return jsonify({'ok': True, 'payload_file': fpath, 'bat_started': started})

if __name__ == '__main__':
    host = os.getenv('HOST', '127.0.0.1')
    try:
        port = int(os.getenv('PORT', '5051'))
    except ValueError:
        port = 5051
    # جلوگیری از اجرا روی 127.0.0.1:5000 به صورت قطعی
    if host.strip() == '127.0.0.1' and port == 5000:
        logger.info("PORT=5000 ممنوع است؛ تغییر به 5051")
        port = 5051
    app.run(host=host, port=port, debug=False, use_reloader=False)
